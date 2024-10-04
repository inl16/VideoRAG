import pandas as pd
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document
import re

# Load the Excel file with video titles and links
file_path = 'LiamOttley.xlsx'
df = pd.read_excel(file_path)

# Create a Word document to save the transcripts
doc = Document()

# Function to extract video ID from a YouTube URL
def extract_video_id(url):
    video_id_match = re.search(r'(v=|\/)([0-9A-Za-z_-]{11}).*', url)
    if video_id_match:
        return video_id_match.group(2)
    return None

# Function to convert seconds to hours:minutes:seconds format
def seconds_to_hms(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = int(seconds % 60)
    return f"{hours:02}:{minutes:02}:{seconds:02}"

# Function to get transcript using youtube-transcript-api
def get_transcript(video_id):
    try:
        # Fetch the transcript
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
        # Combine all transcript lines into a string with time in hh:mm:ss format
        transcript_text = "\n".join([f"Timestamp: {seconds_to_hms(entry['start'])} - Transcript: {entry['text']}" for entry in transcript])
        return transcript_text
    except Exception as e:
        return f"Error fetching transcript: {str(e)}"

# Iterate through the DataFrame, fetch transcripts, and write to Word document
for index, row in df.iterrows():
    video_title = row['Video Title']  # Assuming 'Video Title' is a column in your Excel file
    youtube_url = row['Video Link']   # Assuming 'Video Link' is a column in your Excel file
    
    # Extract video ID from the URL
    video_id = extract_video_id(youtube_url)
    
    if video_id:
        # Add video title and URL to the Word document with designators
        doc.add_heading(f"Video Title: {video_title}", level=1)
        doc.add_paragraph(f"Video Link: {youtube_url}")
        
        # Get the transcript for the video
        transcript = get_transcript(video_id)
        
        # Add transcript to the Word document
        doc.add_paragraph(transcript)
        doc.add_page_break()

# Save the Word document
output_path = 'AAAgency.docx'
doc.save(output_path)
print(f"Transcripts saved to {output_path}")








# import pandas as pd
# from youtube_transcript_api import YouTubeTranscriptApi
# from docx import Document
# import re

# # Load the Excel file with video titles and links
# file_path = 'LiamOttley.xlsx'
# df = pd.read_excel(file_path)

# # Create a Word document to save the transcripts
# doc = Document()

# # Function to extract video ID from a YouTube URL
# def extract_video_id(url):
#     video_id_match = re.search(r'(v=|\/)([0-9A-Za-z_-]{11}).*', url)
#     if video_id_match:
#         return video_id_match.group(2)
#     return None

# # Function to convert seconds to hours:minutes:seconds format
# def seconds_to_hms(seconds):
#     hours = int(seconds // 3600)
#     minutes = int((seconds % 3600) // 60)
#     seconds = int(seconds % 60)
#     return f"{hours:02}:{minutes:02}:{seconds:02}"

# # Function to get transcript using youtube-transcript-api
# def get_transcript(video_id):
#     try:
#         # Fetch the transcript
#         transcript = YouTubeTranscriptApi.get_transcript(video_id)
#         # Combine all transcript lines into a string with time in hh:mm:ss format
#         transcript_text = "\n".join([f"{seconds_to_hms(entry['start'])}: {entry['text']}" for entry in transcript])
#         return transcript_text
#     except Exception as e:
#         return f"Error fetching transcript: {str(e)}"

# # Iterate through the DataFrame, fetch transcripts, and write to Word document
# for index, row in df.iterrows():
#     video_title = row['Video Title']  # Assuming 'Title' is a column in your Excel file
#     youtube_url = row['Video Link']    # Assuming 'URL' is a column in your Excel file
    
#     # Extract video ID from the URL
#     video_id = extract_video_id(youtube_url)
    
#     if video_id:
#         # Add video title and URL to the Word document
#         doc.add_heading(video_title, level=1)
#         doc.add_paragraph(f"URL: {youtube_url}")
        
#         # Get the transcript for the video
#         transcript = get_transcript(video_id)
        
#         # Add transcript to the Word document
#         doc.add_paragraph(transcript)
#         doc.add_page_break()

# # Save the Word document
# output_path = 'AAAgency.docx'
# doc.save(output_path)
# print(f"Transcripts saved to {output_path}")

